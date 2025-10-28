import logging
from typing import Dict, Any
import json
import re
import os
from pathlib import Path

logger = logging.getLogger(__name__)

def parse_pdf(file_path: str) -> Dict[str, Any]:
    """
    Parse PDF resume file and extract text and structured data
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        logger.info(f"Successfully parsed PDF: {file_path}")
        return {"text": text, "pages": len(doc)}
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        return {"text": "", "error": str(e)}

def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Parse DOCX resume file and extract text and structured data
    """
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        logger.info(f"Successfully parsed DOCX: {file_path}")
        return {"text": text, "paragraphs": len(doc.paragraphs)}
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return {"text": "", "error": str(e)}

def extract_contact_info(text: str) -> Dict[str, Any]:
    """
    Extract contact information from resume text
    """
    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    phones = re.findall(r'\b(?:\+?1[-.]?)?$$?([0-9]{3})$$?[-.]?([0-9]{3})[-.]?([0-9]{4})\b', text)
    
    return {
        "emails": emails[:1] if emails else [],  # Get first email
        "phones": [f"{p[0]}-{p[1]}-{p[2]}" for p in phones[:1]] if phones else []  # Get first phone
    }

def extract_skills(text: str) -> list:
    """
    Extract skills from resume text using keyword matching
    """
    common_skills = {
        "programming": ["python", "java", "javascript", "c++", "c#", "ruby", "php", "swift", "kotlin", "go", "rust"],
        "web": ["react", "angular", "vue", "node.js", "express", "django", "flask", "fastapi", "html", "css"],
        "databases": ["sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "cassandra"],
        "cloud": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform"],
        "tools": ["git", "jenkins", "gitlab", "github", "jira", "confluence"],
        "soft_skills": ["leadership", "communication", "teamwork", "problem-solving", "project management"]
    }
    
    text_lower = text.lower()
    found_skills = []
    
    for category, skills in common_skills.items():
        for skill in skills:
            if skill in text_lower:
                found_skills.append(skill.title())
    
    return list(set(found_skills))  # Remove duplicates

def extract_experience(text: str) -> list:
    """
    Extract work experience from resume text
    """
    experience_pattern = r'(?:experience|employment|work history)(.*?)(?:education|skills|$)'
    match = re.search(experience_pattern, text, re.IGNORECASE | re.DOTALL)
    
    if match:
        exp_text = match.group(1)
        # Split by common job entry patterns
        jobs = re.split(r'\n(?=[A-Z])', exp_text)
        return [job.strip() for job in jobs if job.strip()]
    
    return []

def extract_education(text: str) -> list:
    """
    Extract education information from resume text
    """
    education_pattern = r'(?:education|academic)(.*?)(?:experience|skills|$)'
    match = re.search(education_pattern, text, re.IGNORECASE | re.DOTALL)
    
    if match:
        edu_text = match.group(1)
        degrees = re.findall(r'(?:bachelor|master|phd|b\.?s\.?|m\.?s\.?|b\.?a\.?|m\.?a\.?|associate)', edu_text, re.IGNORECASE)
        return degrees
    
    return []

def extract_resume_data(text: str) -> Dict[str, Any]:
    """
    Extract all structured data from resume text
    """
    contact_info = extract_contact_info(text)
    skills = extract_skills(text)
    experience = extract_experience(text)
    education = extract_education(text)
    
    # Extract first line as potential name
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    full_name = lines[0] if lines else "Unknown"
    
    return {
        "full_name": full_name,
        "email": contact_info["emails"][0] if contact_info["emails"] else "",
        "phone": contact_info["phones"][0] if contact_info["phones"] else "",
        "skills": skills,
        "experience": experience,
        "education": education,
        "raw_text": text
    }
